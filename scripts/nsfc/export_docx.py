from __future__ import annotations

import re
from pathlib import Path


HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)$")


def _escape_xml(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&apos;")
    )


def _load_md_files(repo_root: Path) -> list[Path]:
    src = repo_root / "proposal" / "source"
    # Enforce deterministic order via filename prefix.
    return sorted(p for p in src.glob("*.md"))


def _export_docx_minimal(repo_root: Path, out_path: Path) -> None:
    """
    Minimal .docx exporter without external dependencies.

    Produces a valid WordprocessingML package with basic paragraphs/headings.
    Bullets are exported as plain-text "• " prefixed lines (no numbering.xml).
    """
    import zipfile

    out_path.parent.mkdir(parents=True, exist_ok=True)

    paragraphs: list[tuple[str, str | None]] = []
    for md_path in _load_md_files(repo_root):
        text = md_path.read_text(encoding="utf-8")
        for raw_line in text.splitlines():
            line = raw_line.rstrip()
            if not line.strip():
                continue

            line = re.sub(r"\s*\[\[CLAIM:C\d{3,}\]\]\s*", "", line)

            m = HEADING_RE.match(line)
            if m:
                level = len(m.group(1))
                title = m.group(2).strip()
                style = "Heading1" if level == 1 else ("Heading2" if level == 2 else "Heading3")
                paragraphs.append((title, style))
                continue

            if line.lstrip().startswith(("-", "*")):
                paragraphs.append(("• " + line.lstrip()[1:].strip(), None))
                continue

            paragraphs.append((line.strip(), None))

    def p_xml(text: str, style: str | None) -> str:
        style_xml = f'<w:pStyle w:val="{style}"/>' if style else ""
        return (
            "<w:p><w:pPr>"
            f"{style_xml}"
            "</w:pPr>"
            f'<w:r><w:t xml:space="preserve">{_escape_xml(text)}</w:t></w:r>'
            "</w:p>"
        )

    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        "<w:body>"
        + "".join(p_xml(t, s) for t, s in paragraphs)
        + '<w:sectPr><w:pgSz w:w="11906" w:h="16838"/></w:sectPr>'
        + "</w:body></w:document>"
    )

    content_types_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>
"""

    rels_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
"""

    styles_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal">
    <w:name w:val="Normal"/>
    <w:qFormat/>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading1">
    <w:name w:val="heading 1"/>
    <w:basedOn w:val="Normal"/>
    <w:uiPriority w:val="9"/>
    <w:qFormat/>
    <w:pPr><w:keepNext/><w:spacing w:before="240" w:after="120"/></w:pPr>
    <w:rPr><w:b/><w:sz w:val="32"/></w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading2">
    <w:name w:val="heading 2"/>
    <w:basedOn w:val="Normal"/>
    <w:uiPriority w:val="9"/>
    <w:qFormat/>
    <w:pPr><w:keepNext/><w:spacing w:before="200" w:after="100"/></w:pPr>
    <w:rPr><w:b/><w:sz w:val="28"/></w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading3">
    <w:name w:val="heading 3"/>
    <w:basedOn w:val="Normal"/>
    <w:uiPriority w:val="9"/>
    <w:qFormat/>
    <w:pPr><w:keepNext/><w:spacing w:before="180" w:after="90"/></w:pPr>
    <w:rPr><w:b/><w:sz w:val="24"/></w:rPr>
  </w:style>
</w:styles>
"""

    with zipfile.ZipFile(out_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types_xml)
        z.writestr("_rels/.rels", rels_xml)
        z.writestr("word/document.xml", document_xml)
        z.writestr("word/styles.xml", styles_xml)


def export_docx(repo_root: Path, out_path: Path, template_path: Path | None) -> None:
    try:
        from docx import Document  # type: ignore
    except ModuleNotFoundError as e:
        _export_docx_minimal(repo_root=repo_root, out_path=out_path)
        return

    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template_path)) if template_path else Document()

    for md_path in _load_md_files(repo_root):
        text = md_path.read_text(encoding="utf-8")
        for raw_line in text.splitlines():
            line = raw_line.rstrip()
            if not line.strip():
                continue

            m = HEADING_RE.match(line)
            if m:
                level = len(m.group(1))
                title = m.group(2).strip()
                if level == 1:
                    doc.add_heading(title, level=1)
                elif level == 2:
                    doc.add_heading(title, level=2)
                else:
                    doc.add_heading(title, level=3)
                continue

            if line.lstrip().startswith(("-", "*")):
                doc.add_paragraph(line.lstrip()[1:].strip(), style="List Bullet")
                continue

            # Strip claim tags from exported text, but keep content.
            clean = re.sub(r"\s*\[\[CLAIM:C\d{3,}\]\]\s*", "", line)
            doc.add_paragraph(clean)

    doc.save(str(out_path))


__all__ = ["export_docx"]
